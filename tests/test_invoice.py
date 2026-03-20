"""인보이스 서비스 + 엔드포인트 테스트.

QA 에이전트 소유.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from app.services.invoice import AMOUNT_CAP, _cap_amount, generate_invoice


# ── _cap_amount 단위 테스트 ──


class TestCapAmount:
    def test_under_cap(self) -> None:
        assert _cap_amount(8000) == 8000

    def test_exact_cap(self) -> None:
        assert _cap_amount(10000) == 10000

    def test_over_cap(self) -> None:
        assert _cap_amount(13000) == AMOUNT_CAP

    def test_none(self) -> None:
        assert _cap_amount(None) == 0

    def test_zero(self) -> None:
        assert _cap_amount(0) == 0

    def test_negative(self) -> None:
        assert _cap_amount(-5000) == 0


# ── generate_invoice 테스트 ──


class TestGenerateInvoice:
    def test_empty_receipts(self) -> None:
        """영수증 0건 → 합계 0."""
        result = generate_invoice([], 2026, 3)
        assert isinstance(result, io.BytesIO)
        doc = Document(result)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "총 건수: 0건" in text
        assert "₩0" in text

    def test_single_receipt_under_cap(self) -> None:
        """1건, 캡 이하 금액."""
        receipts = [
            {"receipt_date": "2026-03-10", "amount": 8500, "image_path": None},
        ]
        result = generate_invoice(receipts, 2026, 3)
        doc = Document(result)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "2026-03-10" in text
        assert "₩8,500" in text
        assert "총 건수: 1건" in text

    def test_amount_capped(self) -> None:
        """캡 초과 금액 → ₩10,000."""
        receipts = [
            {"receipt_date": "2026-03-05", "amount": 25000, "image_path": None},
        ]
        result = generate_invoice(receipts, 2026, 3)
        doc = Document(result)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "₩10,000" in text
        assert "₩25,000" not in text

    def test_null_date(self) -> None:
        """날짜 없는 영수증 → '날짜 미확인'."""
        receipts = [
            {"receipt_date": None, "amount": 5000, "image_path": None},
        ]
        result = generate_invoice(receipts, 2026, 3)
        doc = Document(result)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "날짜 미확인" in text

    def test_total_uses_capped(self) -> None:
        """합계는 캡 적용된 금액 기준."""
        receipts = [
            {"receipt_date": "2026-03-01", "amount": 15000, "image_path": None},
            {"receipt_date": "2026-03-02", "amount": 8000, "image_path": None},
            {"receipt_date": "2026-03-03", "amount": 20000, "image_path": None},
        ]
        # 캡 적용: 10000 + 8000 + 10000 = 28000
        result = generate_invoice(receipts, 2026, 3)
        doc = Document(result)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "총 건수: 3건" in text
        assert "₩28,000" in text

    def test_sorted_by_date(self) -> None:
        """날짜순 오름차순 정렬."""
        receipts = [
            {"receipt_date": "2026-03-15", "amount": 5000, "image_path": None},
            {"receipt_date": "2026-03-01", "amount": 3000, "image_path": None},
            {"receipt_date": "2026-03-10", "amount": 7000, "image_path": None},
        ]
        result = generate_invoice(receipts, 2026, 3)
        doc = Document(result)
        texts = [p.text for p in doc.paragraphs if "2026-03" in p.text]
        dates = [t.split()[0] for t in texts]
        assert dates == ["2026-03-01", "2026-03-10", "2026-03-15"]

    def test_title_contains_month(self) -> None:
        """제목에 연월 포함."""
        result = generate_invoice([], 2026, 3)
        doc = Document(result)
        # heading은 paragraphs에 포함됨
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "2026년 3월" in text


# ── export 엔드포인트 테스트 ──


@pytest.mark.asyncio
async def test_export_returns_docx(client) -> None:
    """export 엔드포인트 → DOCX 다운로드."""
    resp = await client.get("/api/receipts/export", params={"month": "2026-03"})
    assert resp.status_code == 200
    assert "wordprocessingml" in resp.headers["content-type"]
    assert "attachment" in resp.headers["content-disposition"]
    assert "invoice_2026_03.docx" in resp.headers["content-disposition"]


@pytest.mark.asyncio
async def test_export_invalid_month(client) -> None:
    """잘못된 월 형식 → 400."""
    resp = await client.get("/api/receipts/export", params={"month": "invalid"})
    assert resp.status_code == 400
