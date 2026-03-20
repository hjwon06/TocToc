"""식비 인보이스 DOCX 생성 서비스.

A4 비즈니스 에이전트 소유.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image

logger = logging.getLogger(__name__)

AMOUNT_CAP = 10_000


def _cap_amount(amount: int | None) -> int:
    """금액 캡 적용: 최대 ₩10,000."""
    if amount is None or amount <= 0:
        return 0
    return min(amount, AMOUNT_CAP)


def _convert_image_for_docx(image_path: str) -> io.BytesIO | None:
    """이미지를 DOCX 삽입 가능한 형식(PNG)으로 변환.

    HEIC 등 비호환 포맷도 Pillow로 변환 시도.
    """
    try:
        path = Path(image_path)
        if not path.exists():
            logger.warning("이미지 파일 없음: %s", image_path)
            return None

        with Image.open(path) as img:
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            return buf
    except Exception as e:
        logger.error("이미지 변환 실패: %s — %s", image_path, e)
        return None


def generate_invoice(
    receipts: list[dict],
    year: int,
    month: int,
) -> io.BytesIO:
    """월별 식비 인보이스 DOCX 생성.

    Args:
        receipts: Receipt dict 목록 (receipt_date, amount, image_path 포함).
        year: 연도.
        month: 월.

    Returns:
        DOCX 파일 바이트 스트림.
    """
    doc = Document()

    # ── 기본 스타일 설정 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(11)

    # ── 제목 ──
    title = doc.add_heading(f"식비 인보이스 — {year}년 {month}월", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 날짜순 정렬 (None은 맨 뒤) ──
    sorted_receipts = sorted(
        receipts,
        key=lambda r: r.get("receipt_date") or "9999-99-99",
    )

    total_capped = 0
    count = 0

    for receipt in sorted_receipts:
        raw_date = receipt.get("receipt_date")
        raw_amount = receipt.get("amount")
        image_path = receipt.get("image_path")

        date_text = raw_date if raw_date else "날짜 미확인"
        capped = _cap_amount(raw_amount)
        total_capped += capped
        count += 1

        # 날짜 + 금액 한 줄
        p = doc.add_paragraph()
        run_date = p.add_run(f"{date_text}    ")
        run_date.bold = True
        run_date.font.size = Pt(13)

        run_amount = p.add_run(f"₩{capped:,}")
        run_amount.font.size = Pt(13)

        # 이미지 삽입
        if image_path:
            img_buf = _convert_image_for_docx(image_path)
            if img_buf:
                try:
                    doc.add_picture(img_buf, width=Cm(12))
                except Exception as e:
                    logger.error("이미지 삽입 실패: %s — %s", image_path, e)
                    doc.add_paragraph("(이미지 로드 실패)")
            else:
                doc.add_paragraph("(이미지 없음)")
        else:
            doc.add_paragraph("(이미지 없음)")

        # 구분선
        doc.add_paragraph("─" * 40)

    # ── 합계 ──
    doc.add_paragraph("")
    summary = doc.add_paragraph()
    run_summary = summary.add_run(f"총 건수: {count}건    총 금액: ₩{total_capped:,}")
    run_summary.bold = True
    run_summary.font.size = Pt(14)

    # ── 바이트 스트림으로 반환 ──
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
